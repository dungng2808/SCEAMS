from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path('/Users/dungng/FPT/Ki 8/PRN232/Project_V2/SCEAMS_Phase134_135_Tai_lieu_ky_thuat_va_huong_dan_chay.docx')
BLUE='2E74B5'; NAVY='17365D'; LIGHT='E8EEF5'; PALE='F5F8FC'; INK='1F2937'; MUTED='5B6573'

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill); shd.set(qn('w:val'), 'clear')

def borders(cell, color='D5DFEB'):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el=b.find(qn('w:'+edge))
        if el is None: el=OxmlElement('w:'+edge); b.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6'); el.set(qn('w:color'),color); el.set(qn('w:space'),'0')

def margins(cell):
    tcPr=cell._tc.get_or_add_tcPr(); m=tcPr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tcPr.append(m)
    for name,val in [('top','75'),('start','115'),('bottom','75'),('end','115')]:
        el=OxmlElement('w:'+name); el.set(qn('w:w'),val); el.set(qn('w:type'),'dxa'); m.append(el)

def geometry(table, widths):
    table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    pr=table._tbl.tblPr; w=pr.find(qn('w:tblW'))
    if w is None: w=OxmlElement('w:tblW'); pr.append(w)
    w.set(qn('w:w'),'9360'); w.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); pr.append(ind)
    grid=table._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for width in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(width)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            cell.width=Inches(widths[i]/1440); margins(cell); borders(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcW=cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'),str(widths[i])); tcW.set(qn('w:type'),'dxa')

def repeat_header(row):
    trPr=row._tr.get_or_add_trPr(); x=OxmlElement('w:tblHeader'); x.set(qn('w:val'),'true'); trPr.append(x)

def page_number(p):
    r=p.add_run(); a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); b=OxmlElement('w:instrText'); b.set(qn('xml:space'),'preserve'); b.text='PAGE'; c=OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'),'end'); r._r.extend([a,b,c])

def setup(doc):
    s=doc.sections[0]; s.top_margin=Inches(.82); s.bottom_margin=Inches(.72); s.left_margin=Inches(.82); s.right_margin=Inches(.82); s.header_distance=Inches(.3); s.footer_distance=Inches(.3)
    n=doc.styles['Normal']; n.font.name='Calibri'; n._element.rPr.rFonts.set(qn('w:eastAsia'),'Calibri'); n.font.size=Pt(10.1); n.font.color.rgb=RGBColor.from_string(INK); n.paragraph_format.space_after=Pt(5); n.paragraph_format.line_spacing=1.15
    for name,size,color,before,after in [('Heading 1',16,BLUE,16,8),('Heading 2',13,BLUE,12,5),('Heading 3',11.5,NAVY,9,4)]:
        st=doc.styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Calibri'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for name in ['List Bullet','List Number']:
        st=doc.styles[name]; st.font.name='Calibri'; st.font.size=Pt(10.1); st.paragraph_format.left_indent=Inches(.28); st.paragraph_format.first_line_indent=Inches(-.16); st.paragraph_format.space_after=Pt(3); st.paragraph_format.line_spacing=1.1
    if 'Code Block' not in doc.styles: code=doc.styles.add_style('Code Block',WD_STYLE_TYPE.PARAGRAPH)
    else: code=doc.styles['Code Block']
    code.font.name='Consolas'; code._element.rPr.rFonts.set(qn('w:eastAsia'),'Consolas'); code.font.size=Pt(8.1); code.paragraph_format.left_indent=Inches(.14); code.paragraph_format.right_indent=Inches(.14); code.paragraph_format.space_after=Pt(3); code.paragraph_format.line_spacing=1.0
    if 'Small Note' not in doc.styles: note=doc.styles.add_style('Small Note',WD_STYLE_TYPE.PARAGRAPH)
    else: note=doc.styles['Small Note']
    note.font.name='Calibri'; note.font.size=Pt(8.5); note.font.italic=True; note.font.color.rgb=RGBColor.from_string(MUTED); note.paragraph_format.space_after=Pt(4)
    hp=s.header.paragraphs[0]; hp.text='SCEAMS  |  TECHNICAL SUBMISSION                         Phase 134–135'; hp.runs[0].font.size=Pt(8); hp.runs[0].font.bold=True; hp.runs[0].font.color.rgb=RGBColor.from_string(BLUE)
    fp=s.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; fp.text='SCEAMS • PRN232    |    Trang '; fp.runs[0].font.size=Pt(8); fp.runs[0].font.color.rgb=RGBColor.from_string(MUTED); page_number(fp)

def title(doc,text,sub=''):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(4); r=p.add_run(text); r.font.size=Pt(23); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    if sub: q=doc.add_paragraph(sub); q.paragraph_format.space_after=Pt(11); q.runs[0].font.size=Pt(11.3); q.runs[0].font.color.rgb=RGBColor.from_string(MUTED)

def h(doc,text,level=1): return doc.add_paragraph(text,style='Heading '+str(level))

def bullets(doc,items,numbered=False):
    for item in items: doc.add_paragraph(item,style='List Number' if numbered else 'List Bullet')

def code(doc,text):
    for line in text.strip('\n').splitlines():
        p=doc.add_paragraph(style='Code Block'); pPr=p._p.get_or_add_pPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'F3F6F9'); pPr.append(sh); p.add_run(line)

def callout(doc,head,text,fill=PALE,color=BLUE):
    t=doc.add_table(rows=1,cols=1); geometry(t,[9360]); c=t.cell(0,0); shade(c,fill); borders(c,'C9D9EA'); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); r=p.add_run(head); r.bold=True; r.font.color.rgb=RGBColor.from_string(color); q=c.add_paragraph(text); q.paragraph_format.space_after=Pt(0); q.paragraph_format.line_spacing=1.05; doc.add_paragraph().paragraph_format.space_after=Pt(1)

def table(doc,headers,rows,widths,size=8.0):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; geometry(t,widths); repeat_header(t.rows[0])
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; shade(c,LIGHT); p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0; r=p.add_run(x); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(NAVY)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,x in enumerate(row):
            if ri%2: shade(cells[i],'FAFCFE')
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0; r=p.add_run(str(x)); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(INK)
    geometry(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); return t

def build():
    doc=Document(); setup(doc)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(25); r=p.add_run('SCEAMS'); r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    title(doc,'Tài liệu kỹ thuật bắt buộc\n& Hướng dẫn chạy','Phase 134–135  |  PRN232 • Student Club Event Activity Management System')
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); r=p.add_run('Bản nộp kỹ thuật'); r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY)
    callout(doc,'Mục đích tài liệu','Tập hợp nội dung bắt buộc của Phase 134 và Phase 135 để người mới có thể hiểu kiến trúc, kiểm tra hợp đồng API và chạy SCEAMS từ đầu.','EEF5FB')
    callout(doc,'Phạm vi triển khai','Hệ thống dùng .NET 8, một solution có đúng ba project: SCEAMS.API, SCEAMS.MVC và SCEAMS.NotificationService. API được tách bốn folder/layer Clean Architecture.','F6F8FA',NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.add_run('Ngày lập: ').bold=True; p.add_run('27/07/2026')
    p=doc.add_paragraph(); p.add_run('Trạng thái: ').bold=True; r=p.add_run('Sẵn sàng nộp tài liệu Phase 134–135'); r.font.color.rgb=RGBColor.from_string('2E7D32')
    doc.add_paragraph('Ghi chú: Phần AI FAQ được mô tả như kiến trúc mục tiêu cho Milestone L; không coi đó là tính năng đã bật trong Milestone K.',style='Small Note')
    doc.add_page_break()
    h(doc,'Mục lục nội dung')
    bullets(doc,['Phần I — Phase 134: Tài liệu kỹ thuật bắt buộc','1. Giới thiệu, mục tiêu và phạm vi SCEAMS','2. Actor và use case theo bốn role','3. ERD, khóa, quan hệ và migration/script','4. Business rules và workflow trạng thái','5. Cấu trúc solution và trách nhiệm các layer/service','6. Danh sách endpoint, DTO, role và status code','7. Security matrix','8. OData và content negotiation JSON/XML','9. gRPC service và sequence Web API → NotificationService','10. Kiến trúc AI FAQ, giới hạn và bảo vệ dữ liệu','Phần II — Phase 135: Hướng dẫn chạy','11. Prerequisites và cấu hình','12. Restore, migration, seed và chạy ba project','13. URL, tài khoản demo và checklist xác nhận'])
    doc.add_page_break()
    title(doc,'Phần I — Phase 134','Tài liệu kỹ thuật bắt buộc')
    callout(doc,'Kết quả cần đạt','Người đọc có thể truy vết từ nghiệp vụ → actor/use case → dữ liệu → endpoint → bảo vệ → cách chạy và kiểm thử.','EEF5FB')

    h(doc,'1. Giới thiệu, mục tiêu và phạm vi SCEAMS')
    doc.add_paragraph('SCEAMS (Student Club Event Activity Management System) là hệ thống quản lý hoạt động câu lạc bộ sinh viên: quản lý câu lạc bộ và danh mục, địa điểm, sự kiện, đăng ký, điểm danh, phản hồi, báo cáo và thông báo nhắc lịch. Hệ thống phục vụ quy trình từ tạo bản nháp đến phê duyệt, tổ chức, hoàn tất và tổng hợp số liệu.')
    bullets(doc,['Mục tiêu: chuẩn hóa workflow sự kiện; giảm thao tác thủ công; kiểm soát sức chứa và lịch địa điểm; cung cấp audit/notification; tạo số liệu cho Admin và Staff.','Phạm vi chức năng: JWT/refresh token, RBAC bốn role, CRUD và workflow club, venue, event, registration, attendance, feedback, reports, notification log và health check.','Phạm vi kỹ thuật: .NET 8; SQL Server + EF Core Code First; ASP.NET Core Web API; ASP.NET Core MVC server-rendered; ASP.NET Core gRPC; RFC ProblemDetails; OData trên event list; JSON/XML content negotiation.','Ngoài phạm vi hiện tại: nhà cung cấp AI thật và giao diện chatbot production. AI FAQ ở mục 10 là thiết kế mục tiêu sau Milestone K.'])

    h(doc,'2. Actor và use case theo bốn role')
    table(doc,['Actor','Trách nhiệm chính','Use case tiêu biểu'],[('Admin','Quản trị toàn hệ thống, user, category, venue và phê duyệt.','Quản lý user/role; approve/reject club/event; báo cáo; reminder; notification log; sync status.'),('Staff','Vận hành nghiệp vụ được phân quyền và kiểm duyệt.','Duyệt club/event; quản lý venue; báo cáo; reminder; tra cứu log.'),('Organizer','Đại diện câu lạc bộ tạo và vận hành event.','Tạo club/event draft; submit; sửa/cancel trong ownership; xem registrations; check-in.'),('Student','Người tham gia hoạt động và gửi phản hồi.','Đăng ký/hủy event; xem registrations; join club; feedback sau attended.')],[1450,3300,4610],8.1)
    doc.add_paragraph('Anonymous visitor được xem dữ liệu công khai của event, club, venue, category và feedback; không phải role nghiệp vụ.',style='Small Note')

    h(doc,'3. ERD, khóa, quan hệ và migration/script')
    doc.add_paragraph('ERD logic thể hiện bảng lõi, PK/FK và cardinality. Cột audit (CreatedAt, UpdatedAt, CreatedByUserId, ApprovedByUserId khi có) được lưu UTC.')
    table(doc,['Entity','PK','FK chính','Cardinality / ý nghĩa'],[('Users','Id','—','1–N với Clubs, Events, ClubMemberships, Registrations, Feedbacks; refresh token hash lưu cùng user.'),('ClubCategories','Id','—','1–N với Clubs; Admin quản lý.'),('Clubs','Id','CategoryId → ClubCategories; CreatedByUserId → Users','N–1 category; 1–N memberships/events.'),('ClubMemberships','Id','ClubId → Clubs; UserId → Users','N–1 club và N–1 user; unique logic ClubId + UserId.'),('Venues','Id','—','1–N với Events; có capacity và maintenance.'),('Events','Id','ClubId → Clubs; VenueId → Venues; CreatedByUserId/ApprovedByUserId → Users','N–1 club/venue; 1–N registrations/feedbacks/notification deliveries.'),('Registrations','Id','EventId → Events; StudentId → Users','N–1 event/user; unique EventId + StudentId.'),('Attendances','Id','RegistrationId → Registrations; CheckedInByUserId → Users','1–1 registration.'),('Feedbacks','Id','EventId → Events; StudentId → Users','N–1 event/user; feedback sau Attended.'),('NotificationDeliveries','Id','EventId → Events','N–1 event; unique EventId + NotificationType chống gửi trùng.'),('ChatLogs (future)','Id','StudentId → Users','N–1 student; lưu tối thiểu metadata khi bật AI FAQ.')],[1450,900,3300,3710],7.5)
    callout(doc,'Sơ đồ quan hệ rút gọn','Users 1—N Clubs (creator)  |  ClubCategories 1—N Clubs  |  Clubs 1—N Events  |  Venues 1—N Events  |  Events 1—N Registrations  |  Registrations 1—1 Attendances  |  Events 1—N Feedbacks  |  Events 1—N NotificationDeliveries.','F6F8FA',NAVY)
    h(doc,'3.1 Migration và SQL script',2)
    bullets(doc,['Migration hiện có: InitialCreate; AddRefreshTokenRotation; AddNotificationDeliveries.','Thay đổi schema phải đi qua EF Core migration, review và chạy thử trên development.','Connection string không ghi trực tiếp trong tài liệu hoặc commit; dùng User Secrets/environment variables theo Phase 135.'])
    code(doc,"dotnet ef migrations list --project SCEAMS.API/SCEAMS.API.csproj --startup-project SCEAMS.API/SCEAMS.API.csproj\ndotnet ef database update --project SCEAMS.API/SCEAMS.API.csproj --startup-project SCEAMS.API/SCEAMS.API.csproj\ndotnet ef migrations script --project SCEAMS.API/SCEAMS.API.csproj --startup-project SCEAMS.API/SCEAMS.API.csproj --output docs/sceams.sql")

    h(doc,'4. Business rules và workflow trạng thái')
    table(doc,['Mã','Business rule bắt buộc','Điểm kiểm tra'],[('BR-01','Student chỉ đăng ký event Approved và trước RegistrationDeadline.','Create registration; 400/409.'),('BR-02','Không vượt Capacity; xử lý concurrency để không overbooking.','Transaction/unique constraint.'),('BR-03','Một student không có hai registration active cho cùng event.','Unique EventId + StudentId; 409.'),('BR-04','Không hủy sau mốc khóa hoặc khi event đã hoàn tất.','Registration/event cancel.'),('BR-05','Chỉ Attended mới feedback; một student một feedback/event.','Feedback validation.'),('BR-06','Organizer chỉ sửa/submit/cancel club thuộc ownership.','Ownership scope + 403.'),('BR-07','Approve event kiểm tra venue, capacity và lịch trùng.','Approval command; 409.'),('BR-08','Completed/Cancelled không sửa core information.','Update event; 400/409.'),('BR-09','Membership Active chỉ sau approve.','Membership decision.'),('BR-10*','AI FAQ chỉ trả lời context event Approved, không bịa.','Planned Milestone L guardrail.'),('BR-11*','AI FAQ tối đa 10 câu hỏi/giờ/student; vượt trả 429.','Planned rate limiter.')],[820,5300,3240],7.7)
    doc.add_paragraph('* BR-10 và BR-11 là rule thiết kế mục tiêu Milestone L, chưa bật trong Milestone K.',style='Small Note')
    h(doc,'4.1 Workflow trạng thái',2)
    table(doc,['Đối tượng','Trạng thái','Chuyển trạng thái chính'],[('Club','PendingApproval → Approved / Rejected → Dissolved','Organizer/Admin tạo; Admin/Staff approve, reject hoặc dissolve.'),('Event','Draft → PendingApproval → Approved → Ongoing → Completed; Draft/PendingApproval → Rejected; Approved/Ongoing → Cancelled','Organizer submit; Admin/Staff duyệt; sync theo thời gian; cancel theo quyền.'),('Registration','Pending → Confirmed → Attended; Pending/Confirmed → CancelledByStudent','Student register/cancel; Organizer check-in.'),('ClubMembership','Pending → Active / Rejected → Removed','Student join; Organizer/Admin/Staff quyết định/remove.')],[1450,3300,4610],7.8)

    h(doc,'5. Cấu trúc solution và trách nhiệm các layer/service')
    doc.add_paragraph('SCEAMS.sln có đúng ba project. API là một project duy nhất nhưng tách bốn folder/layer theo Clean Architecture; MVC và gRPC không trộn trách nhiệm với API domain.')
    table(doc,['Project / layer','Trách nhiệm','Ví dụ thành phần'],[('SCEAMS.API / Domain','Entity, enum, invariant và domain rule thuần; không phụ thuộc EF/HTTP.','Event, Club, Registration; EventStatus, UserRole.'),('SCEAMS.API / Application','Use case, DTO, validator, interface và orchestration nghiệp vụ.','EventService, RegistrationService, INotificationClientService, DTOs.'),('SCEAMS.API / Infrastructure','EF Core DbContext/repository, JWT, seed, background worker, gRPC client và provider.','SceamsDbContext, migrations, NotificationClientService, EventReminderBackgroundService.'),('SCEAMS.API / Api','HTTP boundary: controller, model binding, auth policy, middleware, ProblemDetails và Swagger.','EventsController, AuthController, ApiExceptionHandlingMiddleware.'),('SCEAMS.MVC','UI server-rendered; typed HttpClient, token session/cookie flow và lỗi thân thiện.','Controllers, Views, ApiClients, BearerTokenHandler.'),('SCEAMS.NotificationService','gRPC endpoint nhận notification, dedup theo correlation/event và acknowledgement.','notification.proto, NotificationGrpcService.')],[2350,3900,3110],7.8)
    h(doc,'5.1 Sơ đồ solution',2)
    code(doc,"SCEAMS.sln\n├── SCEAMS.API                  # Web API + Clean Architecture folders\n│   ├── Api                    # Controllers, middleware, Swagger\n│   ├── Application            # DTOs, interfaces, use cases\n│   ├── Domain                 # Entities, enums, invariants\n│   └── Infrastructure         # EF Core, JWT, seed, gRPC client, workers\n├── SCEAMS.MVC                 # ASP.NET Core MVC client\n└── SCEAMS.NotificationService # ASP.NET Core gRPC server")

    h(doc,'6. Danh sách endpoint, DTO, role và status code')
    doc.add_paragraph('Quy ước: success dùng 200/201/204; validation 400; token thiếu/sai 401; không đủ quyền 403; không tìm thấy 404; vi phạm nghiệp vụ/duplicate 409; Accept không hỗ trợ 406; lỗi ngoài dự kiến 500. API trả RFC ProblemDetails cho lỗi.')
    h(doc,'6.1 Auth, user, health và category',2)
    table(doc,['Method','Route','DTO / response','Role','Status chính'],[('POST','/api/auth/register','RegisterStudentRequestDto → LoginResponseDto','Anonymous','201, 400, 409'),('POST','/api/auth/login','LoginRequestDto → LoginResponseDto','Anonymous','200, 400, 401'),('POST','/api/auth/refresh','RefreshTokenRequestDto → RefreshTokenResponseDto','Anonymous','200, 401'),('POST','/api/auth/revoke','RefreshTokenRequestDto','Anonymous','204, 400, 401'),('GET','/api/users','PagedUsersResponseDto','Admin','200, 401, 403'),('POST','/api/users','CreateUserRequestDto → CreatedUserResponseDto','Admin','201, 400, 409'),('PUT','/api/users/{id}','UpdateUserRequestDto → UpdatedUserResponseDto','Admin','200, 400, 404'),('PUT','/api/users/{id}/active-status','UpdateUserActiveStatusRequestDto','Admin','204, 400, 404'),('PUT','/api/users/{id}/role','UpdateUserRoleRequestDto','Admin','204, 400, 404'),('GET','/api/users/me','CurrentUserProfileResponseDto','Authenticated','200, 401'),('PUT','/api/users/me','UpdateCurrentUserProfileRequestDto','Authenticated','200, 400, 401'),('PUT','/api/users/me/password','ChangeCurrentUserPasswordRequestDto','Authenticated','204, 400, 401'),('GET','/api/health','HealthResponseDto','Anonymous','200/503'),('GET','/api/health/database','DatabaseHealthResponseDto','Anonymous','200/503'),('GET','/api/club-categories','ClubCategoryResponseDto[]','Anonymous','200'),('POST/PUT/DELETE','/api/club-categories[/{id}]','Create/UpdateClubCategoryRequestDto','Admin','201/204, 400, 404, 409')],[1000,2500,2850,1250,1760],7.0)
    h(doc,'6.2 Club, membership và venue',2)
    table(doc,['Method','Route','DTO / response','Role','Status chính'],[('GET','/api/clubs[/{id}]','ClubResponseDto / ClubDetailResponseDto','Anonymous','200, 404'),('POST','/api/clubs','CreateClubRequestDto → ClubResponseDto','Organizer, Admin','201, 400, 403'),('PUT','/api/clubs/{id}','UpdateClubRequestDto → ClubResponseDto','Organizer, Admin','200, 400, 403, 404'),('PUT','/api/clubs/{id}/approve','—','Admin, Staff','204, 403, 404, 409'),('PUT','/api/clubs/{id}/reject','RejectClubRequestDto','Admin, Staff','204, 400, 404'),('PUT','/api/clubs/{id}/dissolve','—','Admin, Staff','204, 403, 404, 409'),('POST','/api/clubs/{id}/members','—','Student','201/204, 400, 401, 409'),('GET','/api/clubs/{id}/members/pending','ClubMembershipResponseDto[]','Organizer, Admin, Staff','200, 403, 404'),('PUT','/api/clubs/{id}/members/{userId}/decision','DecideClubMembershipRequestDto','Organizer, Admin, Staff','204, 400, 403, 404'),('PUT','/api/clubs/{id}/members/{userId}/remove','RemoveClubMembershipRequestDto','Organizer, Admin, Staff','204, 403, 404'),('GET','/api/venues','VenueResponseDto[]','Anonymous','200'),('GET','/api/venues/{id}/schedule','VenueScheduleResponseDto','Anonymous','200, 404'),('POST','/api/venues','CreateVenueRequestDto → VenueResponseDto','Admin, Staff','201, 400, 403'),('PUT','/api/venues/{id}','UpdateVenueRequestDto → VenueResponseDto','Admin, Staff','200, 400, 404'),('PUT','/api/venues/{id}/maintenance','UpdateVenueMaintenanceRequestDto','Admin, Staff','204, 400, 404'),('DELETE','/api/venues/{id}','—','Admin','204, 403, 404, 409')],[1000,2700,2700,1500,1460],6.9)
    h(doc,'6.3 Event, registration, feedback và reports',2)
    table(doc,['Method','Route','DTO / response','Role','Status chính'],[('GET','/api/events','EventListResponseDto; OData query','Anonymous','200, 406'),('GET','/api/events/{id}','EventDetailResponseDto','Anonymous','200, 404'),('POST','/api/events','CreateEventRequestDto → EventDetailResponseDto','Organizer','201, 400, 403'),('PUT','/api/events/{id}','UpdateEventRequestDto','Organizer, Admin, Staff','204, 400, 403, 404, 409'),('PUT','/api/events/{id}/submit','—','Organizer','204, 400, 403, 404'),('GET','/api/events/pending-approval','EventListResponseDto[]','Admin, Staff','200, 403'),('PUT','/api/events/{id}/approve','—','Admin, Staff','204, 403, 404, 409'),('PUT','/api/events/{id}/reject','RejectEventRequestDto','Admin, Staff','204, 400, 404'),('PUT','/api/events/{id}/cancel','CancelEventRequestDto','Organizer, Admin, Staff','204, 400, 403, 404'),('GET','/api/events/{id}/registrations','RegisteredStudentResponseDto[]','Admin, Organizer','200, 403, 404'),('POST','/api/registrations','CreateRegistrationRequestDto → RegistrationResponseDto','Student','201, 400, 401, 409'),('PUT','/api/registrations/{id}/cancel','—','Student (owner)','204, 400, 403, 404'),('GET','/api/registrations/me','RegistrationResponseDto[]','Student','200, 401'),('PUT','/api/registrations/{id}/check-in','CheckInResponseDto','Organizer','200, 400, 403, 404, 409'),('POST','/api/events/{id}/feedback','CreateFeedbackRequestDto → FeedbackResponseDto','Student','201, 400, 403, 404, 409'),('GET','/api/events/{id}/feedback','FeedbackSummaryResponseDto','Anonymous','200, 404'),('GET','/api/reports/event-summary','Report DTO','Admin, Staff','200, 403'),('GET','/api/reports/club-activity','Report DTO','Admin, Staff, Organizer','200, 403'),('GET','/api/reports/attendance-rate','Report DTO','Admin, Staff, Organizer','200, 403'),('GET','/api/reports/venue-usage','Report DTO','Admin, Staff','200, 403')],[1000,2850,2600,1450,1460],6.7)
    h(doc,'6.4 Notification và reminder',2)
    table(doc,['Method','Route','DTO / response','Role','Status chính'],[('GET','/api/notifications/logs','NotificationDelivery log DTO[]','Admin, Staff','200, 403'),('POST','/api/reminders/run','ReminderRunResultDto','Admin, Staff (Development)','200, 403, 404')],[1000,2700,3000,1500,1160],7.8)

    h(doc,'7. Security matrix')
    table(doc,['Capability','Admin','Staff','Organizer','Student'],[('Xem dữ liệu công khai','✓','✓','✓','✓'),('Auth/refresh/revoke','✓','✓','✓','✓'),('Quản lý user/role','✓','—','—','—'),('Quản lý category','✓','—','—','—'),('Tạo/sửa club','✓','—','✓ (scope)','—'),('Approve/reject/dissolve club','✓','✓','—','—'),('Quản lý venue','✓','✓','—','—'),('Tạo/submit event','—','—','✓ (scope)','—'),('Approve/reject event','✓','✓','—','—'),('Cancel event','✓','✓','✓ (scope)','—'),('Đăng ký/hủy event','—','—','—','✓ (owner)'),('Xem registrations','✓','—','✓ (scope)','—'),('Check-in','—','—','✓ (scope)','—'),('Feedback','—','—','—','✓ (Attended)'),('Reports','✓','✓','✓ (scope)','—'),('Notification log/reminder','✓','✓','—','—')],[2700,1650,1650,1680,1680],7.8)
    bullets(doc,['JWT access token ngắn hạn + refresh token rotation; refresh token lưu hash và revoke khi dùng lại/logout.','Role policy trên controller kết hợp ownership check ở Application; không tin role do client gửi trong body.','HTTPS cho API/MVC/gRPC; token không ghi log; production ProblemDetails không lộ stack trace, connection string hoặc secret.','Validate input, chống overbooking bằng transaction/unique constraint, correlation ID cho notification và audit timestamp UTC.'])

    h(doc,'8. OData và content negotiation JSON/XML')
    doc.add_paragraph('Endpoint demo là GET /api/events. API bật OData query trên event list với giới hạn $top tối đa 50 và chỉ cho phép query option đã whitelist.')
    table(doc,['Demo','Request','Kết quả mong đợi'],[('OData select/filter/order',"GET /api/events?$select=id,title,status&$filter=status eq 'Approved'&$orderby=startAt desc&$top=10",'200; chỉ trả field đã chọn và lọc Approved.'),('OData count/page','GET /api/events?$count=true&$skip=0&$top=20','200; có count và trang đầu; $top > 50 bị validation.'),('JSON','Accept: application/json','200; body JSON theo EventListResponseDto.'),('XML','Accept: application/xml','200; body XML qua XmlSerializerOutputFormatter.'),('Unsupported media','Accept: text/csv','406 Not Acceptable; MVC hiển thị lỗi API thân thiện.')],[1800,4750,2810],7.8)
    code(doc,'curl -k -H "Accept: application/json" "https://localhost:7069/api/events?$select=id,title,status&$top=5"\ncurl -k -H "Accept: application/xml" "https://localhost:7069/api/events?$count=true&$top=5"\ncurl -k -i -H "Accept: text/csv" "https://localhost:7069/api/events"')
    doc.add_paragraph('MVC có trang development System/ContentNegotiation để gửi ba request trên và xem status/content type.',style='Small Note')

    h(doc,'9. gRPC service và sequence Web API → NotificationService')
    doc.add_paragraph('NotificationService là project gRPC riêng. Web API gọi INotificationClientService; Infrastructure hiện thực client từ notification.proto. CorrelationId và EventId + NotificationType hỗ trợ trace và dedup.')
    code(doc,'Student/Organizer action\n        │\n        ▼\nMVC ──HTTP──> SCEAMS.API Controller\n                 │\n                 ▼\n          Application Event/Reminder Service\n                 │ INotificationClientService\n                 ▼\n          Infrastructure gRPC client\n                 │  SendReminder / correlationId\n                 ▼\n      SCEAMS.NotificationService (HTTPS :7001)\n                 │\n                 └── ack { success, correlationId, deduplicated }\n                 ▲\n                 └── API commit/response ──HTTP──> MVC')
    bullets(doc,['Proto: SCEAMS.NotificationService/Protos/notification.proto; server: NotificationGrpcService.','Địa chỉ mặc định NotificationGrpc:Address = https://localhost:7001; timeout 3 giây, retry có giới hạn.','Nếu gRPC down, log correlation ID và trả ProblemDetails phù hợp hoặc lưu trạng thái retry theo use case.'])

    h(doc,'10. Kiến trúc AI FAQ, giới hạn và bảo vệ dữ liệu')
    callout(doc,'Trạng thái triển khai','Đây là kiến trúc mục tiêu cho Milestone L. Milestone K không bật AI API key, không expose chatbot production và không được coi là tính năng đã hoàn thiện.','FFF8E8','B45F06')
    code(doc,'Student → MVC Chat UI → ChatbotController (API)\n                    ↓\n             AIChatService (Application)\n              ↙                    ↘\n     IEventRetrievalService       IAIProvider\n              ↓                    ↓\n      EF/Repository context   Infrastructure AI adapter\n                    ↓\n        Answer grounded only in approved event context')
    bullets(doc,['Retrieval chỉ lấy 5–10 event Approved phù hợp, kèm title/time/venue/capacity đã kiểm chứng; không có trong context thì nói không biết.','IAIProvider nằm Application; provider thật nằm Infrastructure để thay provider mà không kéo SDK vào Domain.','Giới hạn 10 câu hỏi/giờ/student; giới hạn chiều dài; timeout/circuit breaker; 429 + Retry-After; fallback tìm kiếm event.','Không gửi password, JWT, refresh token, email/phone không cần thiết hoặc dữ liệu private; AI_API_KEY chỉ từ User Secrets/environment; không log prompt raw production.','Chống prompt injection bằng system instruction + context delimiters; AI không thực hiện command; mutation vẫn qua endpoint/RBAC chuẩn.','Nếu lưu audit, chỉ lưu metadata tối thiểu với retention được phê duyệt; redaction PII trước khi prompt.'])

    doc.add_page_break()
    title(doc,'Phần II — Phase 135','Hướng dẫn chạy từ clone đến kiểm tra hệ thống')
    callout(doc,'Nguyên tắc an toàn','Không commit connection string thật, JWT signing key hoặc AI API key. Tài liệu dùng placeholder; máy chạy thật nạp secrets qua User Secrets hoặc environment variables.','FFF8E8','B45F06')
    h(doc,'11. Prerequisites và cấu hình')
    bullets(doc,['.NET SDK 8.x (tôn trọng global.json của repository).','SQL Server 2019+ hoặc SQL Server container đang chạy.','Git và terminal; Visual Studio 2022/VS Code là tùy chọn.','dotnet-ef 8.x nếu chạy migration bằng CLI; HTTPS development certificate đã trust.'])
    code(doc,'dotnet --info\ndotnet tool install --global dotnet-ef --version 8.*   # nếu máy chưa có\ndotnet dev-certs https --trust')
    h(doc,'11.1 User Secrets và environment variables',2)
    doc.add_paragraph('Các key bên dưới là tên cấu hình chuẩn. Thay giá trị bằng local; không dán secret thật vào README, appsettings tracked hoặc tài liệu nộp.')
    code(doc,'dotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "ConnectionStrings:DefaultConnection" "Server=localhost;Database=SCEAMS;Trusted_Connection=True;TrustServerCertificate=True"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "Jwt:SigningKey" "<random-secret-at-least-32-chars>"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "Jwt:Issuer" "SCEAMS"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "Jwt:Audience" "SCEAMS.Client"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "SeedData:AdminPassword" "<local-admin-password>"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "SeedData:StaffPassword" "<local-staff-password>"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "SeedData:OrganizerPassword" "<local-organizer-password>"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "SeedData:StudentPassword" "<local-student-password>"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "NotificationGrpc:Address" "https://localhost:7001"\ndotnet user-secrets --project SCEAMS.API/SCEAMS.API.csproj set "AI:ApiKey" "<only-if-Milestone-L-is-enabled>"')
    doc.add_paragraph('Tên environment tương đương: ConnectionStrings__DefaultConnection, Jwt__SigningKey, Jwt__Issuer, Jwt__Audience, NotificationGrpc__Address, AI__ApiKey và SeedData__*.',style='Small Note')

    h(doc,'12. Restore, migration, seed và chạy ba project')
    h(doc,'12.1 Thứ tự thực hiện',2)
    bullets(doc,['Clone repository và chuyển vào thư mục Project_V2.','Restore/build solution để phát hiện thiếu SDK/package.','Apply migration; sau đó seed dữ liệu development (seed chỉ nên chạy ở Development).','Mở ba terminal: NotificationService trước, rồi API, cuối cùng MVC.','Giữ cả ba process chạy trong lúc kiểm thử luồng gRPC và UI.'],True)
    code(doc,'git clone <repo-url>\ncd Project_V2\ndotnet restore SCEAMS.sln\ndotnet build SCEAMS.sln\n\ndotnet ef database update --project SCEAMS.API/SCEAMS.API.csproj --startup-project SCEAMS.API/SCEAMS.API.csproj\ndotnet run --project SCEAMS.API/SCEAMS.API.csproj -- --seed\n\n# Terminal 1 — gRPC\ndotnet run --project SCEAMS.NotificationService/SCEAMS.NotificationService.csproj --launch-profile https\n\n# Terminal 2 — Web API\ndotnet run --project SCEAMS.API/SCEAMS.API.csproj --launch-profile https\n\n# Terminal 3 — MVC\ndotnet run --project SCEAMS.MVC/SCEAMS.MVC.csproj --launch-profile https')
    callout(doc,'Nếu database đã có migration','Không xóa database production để xử lý lỗi local. Kiểm tra connection string, chạy migrations list và xem log; chỉ recreate database development khi đã xác nhận đó là database disposable.','F6F8FA',NAVY)
    h(doc,'12.2 Chạy bằng Visual Studio/VS Code',2)
    doc.add_paragraph('Có thể đặt solution startup thành Multiple projects và chạy cả ba project, nhưng vẫn giữ thứ tự gRPC → API → MVC. CLI ở mục 12.1 là cách tái lập dễ nhất cho người chấm.')

    h(doc,'13. URL, tài khoản demo và checklist xác nhận')
    table(doc,['Thành phần','URL/profile mặc định','Mục đích kiểm tra'],[('Swagger Web API','https://localhost:7069/swagger','Xem OpenAPI, authorize Bearer và thử endpoint.'),('Web API HTTP','http://localhost:5195','Profile HTTP dành cho debug.'),('MVC','https://localhost:7034','Đăng nhập, xem event, đăng ký, check-in/feedback.'),('MVC HTTP','http://localhost:5206','Profile HTTP dành cho debug.'),('NotificationService gRPC','https://localhost:7001','Endpoint gRPC nội bộ; không phải UI browser.')],[2300,2850,4210],8.0)
    h(doc,'13.1 Tài khoản demo',2)
    table(doc,['Role','Email','Mật khẩu'],[('Admin','admin@sceams.edu.vn','Lấy từ SeedData:AdminPassword'),('Staff','staff@sceams.edu.vn','Lấy từ SeedData:StaffPassword'),('Organizer','organizer@sceams.edu.vn','Lấy từ SeedData:OrganizerPassword'),('Student','student@sceams.edu.vn','Lấy từ SeedData:StudentPassword')],[1900,4100,3360],8.2)
    doc.add_paragraph('Mật khẩu không được in trong tài liệu nộp; người chạy tự đặt bằng User Secrets/environment variables trước khi seed.',style='Small Note')
    h(doc,'13.2 Checklist sau khi chạy',2)
    bullets(doc,['GET /api/health trả Healthy và GET /api/health/database xác nhận SQL Server kết nối được.','Swagger mở được, login trả access/refresh token và refresh rotation hoạt động.','MVC mở tại https://localhost:7034; login từng role và không gặp lỗi session/token.','GET /api/events chạy với Accept JSON, Accept XML và text/csv trả 406; thử OData $select/$filter/$top.','Organizer tạo/submit event; Admin/Staff approve; Student register; Organizer check-in; Student feedback sau Attended.','NotificationService chạy trước API; xem notification log/correlation ID khi chạy reminder development.','Kiểm tra git diff: không có secret thật, connection string production hoặc AI API key trong file tracked.'])
    callout(doc,'Tiêu chí “clone và chạy không cần hỏi thêm”','README/repository phải có đúng tên project, profile/URL, lệnh restore–migration–seed–run và tên secret như tài liệu này. Khi đổi port, migration hoặc key cấu hình, cập nhật đồng thời README và tài liệu kỹ thuật.','EEF7EE','2E7D32')
    props=doc.core_properties; props.title='SCEAMS — Phase 134–135: Tài liệu kỹ thuật và hướng dẫn chạy'; props.subject='PRN232 SCEAMS technical submission'; props.author='SCEAMS team'; props.keywords='SCEAMS, PRN232, .NET 8, Clean Architecture, API, MVC, gRPC'
    doc.save(OUT)

if __name__=='__main__':
    build(); print(OUT)
